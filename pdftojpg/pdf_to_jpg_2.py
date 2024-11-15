from pdf2image import convert_from_path
import docx
from docx.shared import Cm
import os
from PIL import Image

from pathlib import Path

doc = docx.Document('Шаблон.docx')
list_of_pfds = os.listdir(path='pdf_files')
path_of_pdf = 'pdf_files'
path_of_jpg = 'jpg_files'

def make_jpg_from_pdf(file):
    '''
    Преобразует PFF в несколько jpg изображений, которые
    сохраняются в корневой директории
    :param file: должен быть PDF
    :return: файлы jpg
    '''
    pages = convert_from_path(file, 500, poppler_path = r"poppler-22.12.0/Library/bin")
    for i, page in enumerate(pages):
        name_of_file = "{}/page_{:03d}.jpg".format(path_of_jpg, i)
        page.save(name_of_file, 'JPEG')

def rotate_img(img):
    '''
    Если ширина изображения больше его высоты - 
    поворачивает изображение влево на 90 градусов
    и сохраняет его под тем же именем
    '''
    im = Image.open(img)
    width = im.size[0] 
    height = im.size[1]
    if width > height:
        im_rotate = im.rotate(90, expand=True)
        im_rotate.save(img, quality=25)
        print('Файл' + img + 'повернут')
    im.close()





for file in list_of_pfds:
    doc.add_heading(text=f'{file[:-4]}', level=2)        # Add header with name as pdf
    doc.add_paragraph(style='Рисунок')
    make_jpg_from_pdf(Path(path_of_pdf, file))
    print('jpeg создан')
    list_of_jpeg = os.listdir(path='jpg_files')
    print(list_of_jpeg)
    for jpg in list_of_jpeg:
        rotate_img(f'{path_of_jpg}/{jpg}')
        doc.add_picture(f'{path_of_jpg}/{jpg}', height=Cm(22))          #Вставляем рисунок в ворд
        doc.add_page_break()
        print(f'jpg{jpg} добавлен')
        os.remove(f'{path_of_jpg}/{jpg}')                               #Удаляем файл изображения из папки
        print(f'jpg{jpg} удален')


doc.save('Приложение.docx')
print('Готово')
